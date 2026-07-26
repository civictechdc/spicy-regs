"""Containment tests for the Office parse ``source.py`` runs behind a process gate.

The Docling adapter is model-free, but it is still a large third-party parser
reading untrusted bytes. ``source.py`` therefore never calls it in this process:
it launches ``python -m spicy_regs.docpipeline.adapters.docling`` in its own
session, hands it one job file inside a temporary directory, and reads back one
strict project-owned JSON record.

These tests hold the gate to exactly what it claims:

* what it enforces — a wall clock, byte caps checked *before* a file is read,
  SIGTERM-then-SIGKILL over the child's whole process group, a
  credential-stripped environment, and a controlled temporary directory; and
* what it does not — CPU, resident memory, archive expansion, temporary disk,
  descendant count, network, and filesystem scope. Those are named as
  unenforced, because an honest receipt beats a limit nothing applies.

Nothing here imports Docling. The real-provider path is covered by one
integration test that skips unless the pinned extra is installed, and runs the
provider in the child process where it belongs.
"""

from __future__ import annotations

import ast
import hashlib
import io
import json
import os
import signal
import sys
import time
from collections.abc import Callable, Sequence
from pathlib import Path
from typing import Any

import pytest

from spicy_regs.docpipeline import source as source_module
from spicy_regs.docpipeline.adapters.docling import (
    ADAPTER_MAPPING_REVISION,
    DOCLING_CORE_PACKAGE,
    DOCLING_CORE_VERSION,
    DOCLING_PACKAGE,
    DOCLING_VERSION,
    installed_package_version,
)
from spicy_regs.docpipeline.source import (
    ENFORCED_LIMITS,
    GATE_CLASSIFICATIONS,
    GATE_COMPLETED,
    GATE_EXIT,
    GATE_EXTRA_UNAVAILABLE,
    GATE_INPUT_OVER_LIMIT,
    GATE_MALFORMED_RESULT,
    GATE_RESULT_OVER_LIMIT,
    GATE_SIGNAL,
    GATE_TIMEOUT,
    PARSER_DERIVED_EVIDENCE,
    OBSERVED_LIMITS,
    UNENFORCED_LIMITS,
    WORKER_MODULE,
    ProcessGateLimits,
    ProcessGateReceipt,
    SourceError,
    contained_environment,
    default_worker_command,
    run_contained_parse,
)

OFFICE_BYTES = b"PK\x03\x04office rendition bytes\n"
SOURCE_NAME = "rule.docx"
MEDIA_TYPE = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

FAST = ProcessGateLimits(wall_timeout_seconds=2.0, terminate_grace_seconds=0.5)


def script_command(script: str) -> Callable[[Path], Sequence[str]]:
    """Run one inline script as the worker, exactly as the gate runs the module."""

    def build(job_path: Path) -> Sequence[str]:
        return (sys.executable, "-I", "-c", script, str(job_path))

    return build


READ_JOB = "import json,sys,pathlib\njob=json.loads(pathlib.Path(sys.argv[1]).read_text())\n"

WORKER_SETUP = (
    READ_JOB
    + r"""
import hashlib
payload = pathlib.Path(job["input_path"]).read_bytes()
policy = {"mapping_revision": """
    + repr(ADAPTER_MAPPING_REVISION)
    + r"""}
policy_digest = hashlib.sha256(
    json.dumps(policy, ensure_ascii=False, sort_keys=True, separators=(",", ":")).encode()
).hexdigest()
parser_id = (
    "docling:" + """
    + repr(DOCLING_VERSION)
    + r""" + ":docling-core:" + """
    + repr(DOCLING_CORE_VERSION)
    + r""" + ":" + """
    + repr(ADAPTER_MAPPING_REVISION)
    + r""" + ":" + policy_digest[:16]
)
offsets = {
    "target": "adapter-parsed-text",
    "unit": "unicode-codepoints",
    "interval": "half-open",
}
call = {
    "provider": "docling",
    "operation": "document-parse",
    "parser_id": parser_id,
    "policy": policy,
    "policy_digest": policy_digest,
    "status": "completed",
    "input_format": "docx",
    "source_sha256": hashlib.sha256(payload).hexdigest(),
    "source_bytes": len(payload),
    "source_name_sha256": hashlib.sha256(job["source_name"].encode()).hexdigest(),
    "media_type": job.get("media_type"),
    "evidence_grade": "parser-derived",
    "offsets": offsets,
}
"""
)

COMPLETED_WORKER = (
    WORKER_SETUP
    + r"""
text = "Effluent Guidelines\n\nFacilities must sample water."
title, body = text.split("\n\n", 1)
call.update({
    "element_count": 2,
    "usable_element_count": 2,
    "usable_character_count": len(title) + len(body),
    "character_count": len(text),
    "content_layers_present": ["body"],
    "coordinate_grade": "none",
})
record = {
    "status": "completed",
    "call": call,
    "document": {
        "text": text,
        "input_format": "docx",
        "source_sha256": hashlib.sha256(payload).hexdigest(),
        "source_bytes": len(payload),
        "evidence_grade": "parser-derived",
        "offsets": offsets,
        "elements": [
            {
                "ordinal": 0, "kind": "title", "text": title, "start_char": 0,
                "end_char": len(title), "content_layer": "body",
                "coordinate_grade": "none", "text_usable": True, "heading_path": [],
            },
            {
                "ordinal": 1, "kind": "text", "text": body, "start_char": len(title) + 2,
                "end_char": len(text), "content_layer": "body",
                "coordinate_grade": "none", "text_usable": True, "heading_path": [title],
            },
        ],
    },
}
pathlib.Path(job["result_path"]).write_text(json.dumps(record))
"""
)

FAILED_WORKER = (
    WORKER_SETUP
    + r"""
call.update({"status": "failed", "failure_reason": "no_usable_text"})
pathlib.Path(job["result_path"]).write_text(json.dumps({
    "status": "failed",
    "failure_reason": "no_usable_text",
    "call": call,
}))
"""
)


def gate(script: str, *, limits: ProcessGateLimits = FAST, **options: Any) -> Any:
    return run_contained_parse(
        OFFICE_BYTES,
        source_name=SOURCE_NAME,
        media_type=MEDIA_TYPE,
        limits=limits,
        worker_command=script_command(script),
        **options,
    )


# --- the launched command ---------------------------------------------------


def test_the_default_command_launches_the_adapter_module_in_isolated_mode(tmp_path: Path) -> None:
    command = default_worker_command(tmp_path / "job.json")

    assert command[0] == sys.executable
    assert "-I" in command
    assert command[-2:] == ("-m", WORKER_MODULE) or list(command[-3:]) == [
        "-m",
        WORKER_MODULE,
        str(tmp_path / "job.json"),
    ]
    assert WORKER_MODULE == "spicy_regs.docpipeline.adapters.docling"


def test_the_child_environment_is_an_allowlist_and_carries_no_credential(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    monkeypatch.setenv("OPENAI_API_KEY", "sk-" + "A" * 40)
    monkeypatch.setenv("R2_PUBLIC_URL", "https://example.invalid")
    monkeypatch.setenv("AWS_SECRET_ACCESS_KEY", "secret")

    built = contained_environment()

    assert set(built) <= set(source_module.ENVIRONMENT_ALLOWLIST)
    assert "OPENAI_API_KEY" not in built
    assert "R2_PUBLIC_URL" not in built
    assert "AWS_SECRET_ACCESS_KEY" not in built


def test_the_child_really_receives_no_credential_from_this_process(
    monkeypatch: pytest.MonkeyPatch,
    tmp_path: Path,
) -> None:
    monkeypatch.setenv("OPENAI_API_KEY", "sk-" + "B" * 40)
    seen = tmp_path / "environment.json"
    script = READ_JOB + (
        "import os, json, pathlib\n"
        f"pathlib.Path({str(seen)!r}).write_text(json.dumps(sorted(os.environ)))\n"
        'pathlib.Path(job["result_path"]).write_text(json.dumps({"status": "failed"}))\n'
    )

    gate(script)

    names = json.loads(seen.read_text())
    assert "OPENAI_API_KEY" not in names
    assert not [name for name in names if "KEY" in name or "SECRET" in name or "TOKEN" in name]


def test_the_child_runs_in_a_controlled_temporary_directory_that_is_removed(tmp_path: Path) -> None:
    seen = tmp_path / "cwd.txt"
    script = READ_JOB + (
        "import os, pathlib, json\n"
        f"pathlib.Path({str(seen)!r}).write_text(os.getcwd())\n"
        'pathlib.Path(job["result_path"]).write_text(json.dumps({"status": "failed"}))\n'
    )

    gate(script)

    directory = Path(seen.read_text())
    assert directory.name.startswith("spicy-regs-source-gate-")
    assert not directory.exists()


# --- success ----------------------------------------------------------------


def test_an_office_success_round_trips_through_the_gate_as_project_records() -> None:
    result = gate(COMPLETED_WORKER)

    assert result.receipt.classification == GATE_COMPLETED
    assert result.receipt.parser_status == "completed"
    assert result.receipt.exit_status == 0
    assert result.parsed is not None
    assert result.parsed.text == "Effluent Guidelines\n\nFacilities must sample water."
    assert [element.kind for element in result.parsed.elements] == ["title", "text"]
    assert result.parsed.elements[1].heading_path == ("Effluent Guidelines",)
    assert result.parsed.evidence_grade == PARSER_DERIVED_EVIDENCE
    assert result.parsed.offsets.target == "adapter-parsed-text"
    assert result.receipt.result_sha256


def test_the_worker_result_carries_only_project_owned_json_builtins(tmp_path: Path) -> None:
    copy = tmp_path / "result.json"
    script = (
        COMPLETED_WORKER + f"\npathlib.Path({str(copy)!r}).write_text(pathlib.Path(job['result_path']).read_text())\n"
    )

    gate(script)

    record = json.loads(copy.read_text())
    assert_only_json_builtins(record)


def assert_only_json_builtins(value: Any, path: str = "result") -> None:
    if isinstance(value, dict):
        for key, item in value.items():
            assert type(key) is str, f"{path} key {key!r}"
            assert_only_json_builtins(item, f"{path}.{key}")
        return
    if isinstance(value, list):
        for index, item in enumerate(value):
            assert_only_json_builtins(item, f"{path}[{index}]")
        return
    assert value is None or type(value) in (str, int, float, bool), f"{path} is {type(value).__name__}"


# --- timeout and process-group termination ---------------------------------


def test_a_child_that_never_finishes_is_stopped_and_classified_as_a_timeout() -> None:
    started = time.monotonic()

    result = gate(
        "import time\ntime.sleep(30)\n", limits=ProcessGateLimits(wall_timeout_seconds=0.5, terminate_grace_seconds=0.3)
    )

    assert result.receipt.classification == GATE_TIMEOUT
    assert result.parsed is None
    assert result.receipt.process_group_terminated is True
    assert time.monotonic() - started < 15


def alive(pid: int) -> bool:
    try:
        os.kill(pid, 0)
    except (ProcessLookupError, PermissionError):
        return False
    return True


def wait_until_gone(pid: int, *, seconds: float = 10.0) -> bool:
    deadline = time.monotonic() + seconds
    while time.monotonic() < deadline:
        if not alive(pid):
            return True
        time.sleep(0.05)
    return not alive(pid)


GRANDCHILD = (
    "import subprocess, sys, pathlib, time\n"
    "child = subprocess.Popen([sys.executable, '-I', '-c', 'import time; time.sleep(120)'])\n"
    "pathlib.Path({marker!r}).write_text(str(child.pid))\n"
)


def test_a_timeout_kills_the_whole_process_group_including_a_grandchild(tmp_path: Path) -> None:
    marker = tmp_path / "grandchild.pid"
    script = GRANDCHILD.format(marker=str(marker)) + "time.sleep(120)\n"

    result = gate(script, limits=ProcessGateLimits(wall_timeout_seconds=1.5, terminate_grace_seconds=0.3))

    assert result.receipt.classification == GATE_TIMEOUT
    grandchild = int(marker.read_text())
    assert wait_until_gone(grandchild), "a grandchild survived the process-group kill"


def test_a_clean_exit_still_reaps_a_grandchild_the_child_left_behind(tmp_path: Path) -> None:
    marker = tmp_path / "grandchild.pid"
    script = (
        READ_JOB
        + GRANDCHILD.format(marker=str(marker))
        + ('import json\npathlib.Path(job["result_path"]).write_text(json.dumps({"status": "failed"}))\n')
    )

    result = gate(script)

    assert result.receipt.classification == GATE_MALFORMED_RESULT
    assert result.receipt.process_group_terminated is True
    grandchild = int(marker.read_text())
    assert wait_until_gone(grandchild), "a grandchild outlived a clean gate exit"


def test_a_child_ignoring_sigterm_is_killed_anyway() -> None:
    script = "import signal, time\nsignal.signal(signal.SIGTERM, signal.SIG_IGN)\ntime.sleep(120)\n"

    result = gate(script, limits=ProcessGateLimits(wall_timeout_seconds=1.0, terminate_grace_seconds=0.3))

    assert result.receipt.classification == GATE_TIMEOUT
    assert result.receipt.signal_number == int(signal.SIGKILL)


# --- bounded output ---------------------------------------------------------


def test_a_large_result_is_measured_and_refused_without_being_read(tmp_path: Path) -> None:
    script = READ_JOB + ('import pathlib\npathlib.Path(job["result_path"]).write_bytes(b"{" + b"x" * 200_000)\n')

    result = gate(script, limits=ProcessGateLimits(wall_timeout_seconds=5.0, max_result_bytes=1024))

    assert result.receipt.classification == GATE_RESULT_OVER_LIMIT
    assert result.receipt.result_over_limit is True
    assert result.receipt.result_bytes > 1024
    assert result.parsed is None
    # Nothing was read, so nothing was hashed either.
    assert result.receipt.result_sha256 == ""


def test_the_size_check_happens_before_the_read(tmp_path: Path) -> None:
    """A probe: an unreadable oversized result must classify, not raise.

    If the gate read the file first, this would raise ``PermissionError``
    instead of settling on the recorded byte cap.
    """
    script = READ_JOB + (
        "import pathlib, os\n"
        'target = pathlib.Path(job["result_path"])\n'
        'target.write_bytes(b"x" * 50_000)\n'
        "os.chmod(target, 0o000)\n"
    )

    result = gate(script, limits=ProcessGateLimits(wall_timeout_seconds=5.0, max_result_bytes=1024))

    assert result.receipt.classification == GATE_RESULT_OVER_LIMIT
    assert result.parsed is None


def test_a_child_flooding_stderr_and_stdout_never_deadlocks_the_gate() -> None:
    script = READ_JOB + (
        "import sys, json, pathlib\n"
        "block = 'e' * 100_000\n"
        "for _ in range(60):\n"
        "    sys.stderr.write(block)\n"
        "    sys.stdout.write(block)\n"
        "sys.stderr.flush()\n"
        "sys.stdout.flush()\n" + COMPLETED_WORKER.split(READ_JOB, 1)[1]
    )
    started = time.monotonic()

    result = gate(script, limits=ProcessGateLimits(wall_timeout_seconds=30.0, max_stderr_bytes=1024))

    assert result.parsed is not None, "the gate deadlocked or timed out on child output"
    assert result.receipt.classification == GATE_COMPLETED
    assert result.receipt.stderr_bytes >= 6_000_000
    assert result.receipt.stderr_over_limit is True
    assert time.monotonic() - started < 30


def test_no_receipt_field_repeats_child_output_or_source_text() -> None:
    secret = "sk-" + "C" * 40
    script = WORKER_SETUP + (
        f"sys.stderr.write({secret!r} + ' /var/secrets/scan.docx')\n"
        'call.update({"status": "failed", "failure_reason": "no_usable_text"})\n'
        'pathlib.Path(job["result_path"]).write_text(json.dumps({'
        '"status": "failed", "failure_reason": "no_usable_text", "call": call}))\n'
    )

    result = gate(script)

    serialized = json.dumps(result.receipt.as_dict())
    assert secret not in serialized
    assert "scan.docx" not in serialized
    assert "office rendition bytes" not in serialized
    assert result.receipt.parser_failure_reason == "no_usable_text"


# --- fixed classifications ---------------------------------------------------


def test_a_nonzero_exit_is_recorded_as_such() -> None:
    result = gate("import sys\nsys.exit(3)\n")

    assert result.receipt.classification == GATE_EXIT
    assert result.receipt.exit_status == 3
    assert result.receipt.signal_number is None
    assert result.parsed is None


def test_a_child_killed_by_a_signal_is_recorded_as_such() -> None:
    result = gate("import os, signal\nos.kill(os.getpid(), signal.SIGKILL)\n")

    assert result.receipt.classification == GATE_SIGNAL
    assert result.receipt.signal_number == int(signal.SIGKILL)
    assert result.receipt.exit_status is None


@pytest.mark.parametrize(
    "body",
    [
        "not json at all",
        '{"status": "surprising"}',
        '["a", "list"]',
        '{"status": "completed", "call": {}, "document": {"text": 5}}',
    ],
)
def test_malformed_worker_data_is_one_fixed_classification(body: str) -> None:
    script = READ_JOB + f'import pathlib\npathlib.Path(job["result_path"]).write_text({body!r})\n'

    result = gate(script)

    assert result.receipt.classification == GATE_MALFORMED_RESULT
    assert result.parsed is None


def test_a_missing_result_file_is_malformed_rather_than_a_success() -> None:
    result = gate("pass\n")

    assert result.receipt.classification == GATE_MALFORMED_RESULT
    assert result.parsed is None


def test_a_declared_parse_failure_is_a_completed_gate_with_a_failed_parse() -> None:
    script = FAILED_WORKER.replace("no_usable_text", "format_mismatch")

    result = gate(script)

    assert result.receipt.classification == GATE_COMPLETED
    assert result.receipt.parser_status == "failed"
    assert result.receipt.parser_failure_reason == "format_mismatch"
    assert result.parsed is None
    assert result.call is not None
    assert result.call["status"] == "failed"
    assert result.call["failure_reason"] == "format_mismatch"


def test_every_classification_this_gate_may_record_is_declared() -> None:
    assert set(GATE_CLASSIFICATIONS) == {
        GATE_COMPLETED,
        GATE_TIMEOUT,
        GATE_EXIT,
        GATE_SIGNAL,
        GATE_MALFORMED_RESULT,
        GATE_RESULT_OVER_LIMIT,
        GATE_INPUT_OVER_LIMIT,
        GATE_EXTRA_UNAVAILABLE,
    }
    with pytest.raises(SourceError, match="unknown gate classification"):
        ProcessGateReceipt(
            worker_module=WORKER_MODULE,
            classification="invented",
            parser_status="",
            parser_failure_reason=None,
            exit_status=0,
            signal_number=None,
            process_group_terminated=False,
            duration_ms=0.0,
            result_bytes=0,
            result_over_limit=False,
            stderr_bytes=0,
            stderr_over_limit=False,
            limits=FAST,
        )


# --- exact limit claims -----------------------------------------------------


def test_the_gate_claims_exactly_the_limits_it_enforces() -> None:
    assert ENFORCED_LIMITS == (
        "input_bytes",
        "wall_timeout_seconds",
        "result_bytes",
        "process_group_termination",
        "credential_stripped_environment",
        "controlled_temporary_directory",
    )
    assert UNENFORCED_LIMITS == (
        "cpu_seconds",
        "resident_memory_bytes",
        "archive_expansion_bytes",
        "temporary_disk_bytes",
        "descendant_process_count",
        "network_access",
        "filesystem_scope",
    )
    assert OBSERVED_LIMITS == ("stderr_bytes",)
    assert not set(ENFORCED_LIMITS) & set(OBSERVED_LIMITS)
    assert not set(ENFORCED_LIMITS) & set(UNENFORCED_LIMITS)
    assert not set(OBSERVED_LIMITS) & set(UNENFORCED_LIMITS)


def test_every_receipt_repeats_both_limit_lists_unchanged() -> None:
    result = gate(COMPLETED_WORKER)

    recorded = result.receipt.as_dict()
    assert recorded["enforced_limits"] == list(ENFORCED_LIMITS)
    assert recorded["observed_limits"] == list(OBSERVED_LIMITS)
    assert recorded["unenforced_limits"] == list(UNENFORCED_LIMITS)
    assert recorded["limits"]["wall_timeout_seconds"] == FAST.wall_timeout_seconds


def test_probe_an_unenforced_limit_cannot_be_claimed_as_enforced() -> None:
    with pytest.raises(SourceError, match="enforced, observed, and unenforced"):
        ProcessGateReceipt(
            worker_module=WORKER_MODULE,
            classification=GATE_COMPLETED,
            parser_status="completed",
            parser_failure_reason=None,
            exit_status=0,
            signal_number=None,
            process_group_terminated=True,
            duration_ms=0.0,
            result_bytes=0,
            result_over_limit=False,
            stderr_bytes=0,
            stderr_over_limit=False,
            limits=FAST,
            enforced_limits=(*ENFORCED_LIMITS, "network_access"),
        )


def test_the_gate_refuses_input_bytes_over_its_recorded_limit() -> None:
    result = run_contained_parse(
        b"x" * 100,
        source_name=SOURCE_NAME,
        limits=ProcessGateLimits(max_input_bytes=10),
        worker_command=script_command("pass\n"),
    )

    assert result.receipt.classification == GATE_INPUT_OVER_LIMIT
    assert result.parsed is None


def test_the_gate_never_claims_containment_the_module_does_not_implement() -> None:
    text = Path(source_module.__file__ or "").read_text(encoding="utf-8")
    tree = ast.parse(text)
    called = {
        node.func.attr for node in ast.walk(tree) if isinstance(node, ast.Call) and isinstance(node.func, ast.Attribute)
    }

    # Nothing below sets a resource limit, so nothing above may name one.
    assert "setrlimit" not in called
    assert "resource" not in {
        alias.name for node in ast.walk(tree) if isinstance(node, ast.Import) for alias in node.names
    }
    for name in UNENFORCED_LIMITS:
        assert name in text  # named, and named only in the unenforced list


# --- the worker entry, without a process ------------------------------------


class StubParser:
    """Stands in for the pinned parser so the entry point runs without the extra.

    It satisfies the adapter's ``DocumentParser`` protocol exactly, so this test
    proves the worker holds to that interface rather than to one class.
    """

    provider = "docling"
    parser_id = "docling:2.115.0:docling-core:2.87.1:office-mapping-6:abcdef0123456789"
    production_provider = False
    supported_formats = frozenset({"docx", "pptx", "xlsx"})

    def __init__(self, *, result: Any = None, error: BaseException | None = None, **options: Any) -> None:
        self.result = result
        self.error = error
        self.options = options

    def parse(self, content: bytes, *, source_name: str, media_type: str | None = None) -> Any:
        if self.error is not None:
            raise self.error
        return self.result


def write_job(tmp_path: Path, **changes: Any) -> Path:
    job = {
        "input_path": str(tmp_path / "input.bin"),
        "result_path": str(tmp_path / "result.json"),
        "source_name": SOURCE_NAME,
        "media_type": MEDIA_TYPE,
        "max_source_bytes": 1024,
    }
    job.update(changes)
    (tmp_path / "input.bin").write_bytes(OFFICE_BYTES)
    path = tmp_path / "job.json"
    path.write_text(json.dumps(job), encoding="utf-8")
    return path


def test_the_worker_writes_a_completed_record_for_a_successful_parse(tmp_path: Path) -> None:
    from spicy_regs.docpipeline.adapters.docling import (
        OffsetSemantics,
        ParsedDocument,
        ParsedDocumentResult,
        ParsedElement,
        ParserCall,
        run_worker,
    )

    element = ParsedElement(
        ordinal=0,
        kind="text",
        text="Body.",
        start_char=0,
        end_char=5,
        tree_level=1,
        content_source="provider-text",
        content_layer="body",
        text_usable=True,
        coordinate_grade="none",
    )
    document = ParsedDocument(
        text="Body.",
        elements=(element,),
        tables=(),
        omissions=(),
        source_sha256=hashlib.sha256(OFFICE_BYTES).hexdigest(),
        source_bytes=len(OFFICE_BYTES),
        input_format="docx",
    )
    call = _fake_call(ParserCall, OffsetSemantics)
    job = write_job(tmp_path)

    status = run_worker(
        job,
        parser_factory=lambda **options: StubParser(result=ParsedDocumentResult(document=document, call=call)),
    )

    record = json.loads((tmp_path / "result.json").read_text())
    assert status == 0
    assert record["status"] == "completed"
    assert record["document"]["text"] == "Body."
    assert record["document"]["elements"][0]["heading_path"] == []
    assert record["call"]["parser_id"] == call.parser_id
    assert_only_json_builtins(record)


def _fake_call(parser_call_type: Any, offsets_type: Any) -> Any:
    """One complete adapter call record, built without the optional extra."""
    from spicy_regs.docpipeline.adapters import docling as adapter

    policy = adapter.ParserPolicy(
        pipeline="simple",
        mapping_revision=adapter.ADAPTER_MAPPING_REVISION,
        content_layers=adapter.CONTENT_LAYERS,
        element_separator=adapter.ELEMENT_SEPARATOR,
        heading_kinds=tuple(sorted(adapter.HEADING_KINDS)),
        table_serialization=adapter.TABLE_SERIALIZATION,
        text_offsets=adapter.PARSED_TEXT_OFFSETS,
        supported_formats=tuple(sorted(adapter.SUPPORTED_FORMATS)),
        max_source_bytes=1024,
        max_items=1,
        max_tables=1,
        max_cells_per_table=1,
        max_total_table_cells=1,
        max_table_dimension=1,
        max_mapped_characters=1,
        max_table_cell_characters=1,
        max_heading_level=1,
        max_tree_depth=1,
        max_reference_chars=1,
        max_caption_refs_per_item=1,
        max_total_caption_refs=1,
        max_regions_per_item=1,
        max_total_regions=1,
        max_page_number=1,
        max_provenance_char_index=1,
        max_provider_errors=1,
        max_error_type_chars=1,
        remote_services_enabled=False,
        external_plugins_allowed=False,
        picture_classification_enabled=False,
        picture_description_enabled=False,
        chart_extraction_enabled=False,
        document_timeout_enforced=False,
        page_limit_enforced=False,
        process_containment_enforced=False,
        converter_source="injected",
    )
    return parser_call_type(
        provider="docling",
        operation="document-parse",
        package_name=DOCLING_PACKAGE,
        package_version=DOCLING_VERSION,
        core_package_name=DOCLING_CORE_PACKAGE,
        core_package_version=DOCLING_CORE_VERSION,
        parser_id="docling:2.115.0:docling-core:2.87.1:office-mapping-6:abcdef0123456789",
        policy_digest=policy.digest,
        policy=policy,
        source_name=SOURCE_NAME,
        source_name_sha256=hashlib.sha256(SOURCE_NAME.encode()).hexdigest(),
        source_name_sanitized=False,
        media_type=MEDIA_TYPE,
        input_format="docx",
        source_sha256=hashlib.sha256(OFFICE_BYTES).hexdigest(),
        source_bytes=len(OFFICE_BYTES),
        conversion_status="success",
        provider_error_count=0,
        provider_error_categories=(),
        provider_input_format="docx",
        page_count=None,
        element_count=1,
        usable_element_count=1,
        usable_character_count=5,
        character_count=5,
        content_layers_present=("body",),
        table_count=0,
        table_cell_count=0,
        omission_count=0,
        omitted_kinds=(),
        elements_without_coordinates=1,
        coordinate_grade="none",
        evidence_grade="parser-derived",
        offsets=offsets_type(target="adapter-parsed-text", unit="unicode-codepoints", interval="half-open"),
        status="completed",
        provider_invoked=True,
        attempt_count=1,
        duration_ms=1.0,
        failure_reason=None,
        error_type=None,
    )


def test_the_worker_reports_an_absent_extra_in_one_word(tmp_path: Path) -> None:
    from spicy_regs.docpipeline.adapters.docling import DoclingUnavailableError, run_worker

    def factory(**options: Any) -> Any:
        raise DoclingUnavailableError("docling version differs from the pinned contract")

    status = run_worker(write_job(tmp_path), parser_factory=factory)

    record = json.loads((tmp_path / "result.json").read_text())
    assert status == 0
    assert record == {"status": "unavailable"}


def test_the_worker_records_a_parse_failure_without_repeating_its_message(tmp_path: Path) -> None:
    from spicy_regs.docpipeline.adapters.docling import DoclingParseError, OffsetSemantics, ParserCall, run_worker

    call = _fake_call(ParserCall, OffsetSemantics)
    failure = DoclingParseError("boom /var/secrets/scan.docx", call=call)
    status = run_worker(write_job(tmp_path), parser_factory=lambda **options: StubParser(error=failure))

    record = json.loads((tmp_path / "result.json").read_text())
    assert status == 0
    assert record["status"] == "failed"
    assert record["call"]["parser_id"] == call.parser_id
    assert "scan.docx" not in json.dumps(record)
    assert_only_json_builtins(record)


def test_an_unreadable_job_file_exits_non_zero_and_writes_nothing(tmp_path: Path) -> None:
    from spicy_regs.docpipeline.adapters.docling import run_worker

    status = run_worker(tmp_path / "absent.json", parser_factory=lambda **options: StubParser())

    assert status == 1
    assert not (tmp_path / "result.json").exists()


def test_the_worker_entry_takes_exactly_one_argument() -> None:
    from spicy_regs.docpipeline.adapters.docling import main

    assert main([]) == 2
    assert main(["a", "b"]) == 2


# --- the real adapter module -------------------------------------------------


def test_the_worker_module_answers_a_job_even_with_the_extra_absent(tmp_path: Path) -> None:
    """The real module runs; without the pinned extra it says so, in one word."""
    if installed_package_version(DOCLING_PACKAGE) == DOCLING_VERSION:
        pytest.skip("the pinned docling extra is installed; the unavailable path needs it absent")

    result = run_contained_parse(
        OFFICE_BYTES,
        source_name=SOURCE_NAME,
        media_type=MEDIA_TYPE,
        limits=ProcessGateLimits(wall_timeout_seconds=60.0),
    )

    assert result.receipt.classification == GATE_EXTRA_UNAVAILABLE
    assert result.receipt.worker_module == WORKER_MODULE
    assert result.parsed is None


pinned_extra = pytest.mark.skipif(
    installed_package_version(DOCLING_PACKAGE) != DOCLING_VERSION
    or installed_package_version(DOCLING_CORE_PACKAGE) != DOCLING_CORE_VERSION,
    reason="needs the pinned docling extra: uv run --frozen --extra docling pytest",
)


def headings_docx() -> bytes:
    import docx  # ty: ignore[unresolved-import]

    document = docx.Document()
    document.add_heading("Scope", level=1)
    document.add_paragraph("This part applies to discharges.")
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


@pinned_extra
def test_the_real_adapter_parses_a_real_docx_inside_the_gate() -> None:
    payload = headings_docx()

    result = run_contained_parse(
        payload,
        source_name="scope.docx",
        media_type=MEDIA_TYPE,
        limits=ProcessGateLimits(wall_timeout_seconds=300.0),
    )

    assert result.receipt.classification == GATE_COMPLETED, result.receipt.as_dict()
    assert result.parsed is not None
    assert "This part applies to discharges." in result.parsed.text
    assert result.parsed.input_format == "docx"
    assert result.parsed.source_sha256 == hashlib.sha256(payload).hexdigest()
    assert result.parsed.evidence_grade == PARSER_DERIVED_EVIDENCE
    assert result.parsed.parser_id.startswith(f"{DOCLING_PACKAGE}:{DOCLING_VERSION}")
    for element in result.parsed.elements:
        assert result.parsed.text[element.start_char : element.end_char] == element.text


@pinned_extra
def test_a_real_pdf_is_refused_by_the_real_worker_without_a_paginated_parse() -> None:
    result = run_contained_parse(
        b"%PDF-1.7\n%\xe2\xe3\xcf\xd3\n",
        source_name="rule.pdf",
        media_type="application/pdf",
        limits=ProcessGateLimits(wall_timeout_seconds=120.0),
    )

    assert result.receipt.classification == GATE_COMPLETED
    assert result.receipt.parser_status == "failed"
    assert result.receipt.parser_failure_reason == "format_not_implemented"
    assert result.parsed is None
